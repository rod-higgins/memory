"""Document data source adapters (Notion, Obsidian, Notes, etc.)."""

from __future__ import annotations

import sqlite3
from collections.abc import AsyncIterator
from datetime import datetime
from pathlib import Path

from .base import DataCategory, DataPoint, DataSource


class ObsidianVaultSource(DataSource):
    """Adapter for Obsidian vault (markdown files)."""

    source_type = "obsidian"
    category = DataCategory.NOTES

    def __init__(
        self,
        vault_path: str,
        include_daily_notes: bool = True,
        include_templates: bool = False,
        **kwargs,
    ):
        """
        Initialize Obsidian vault source.

        Args:
            vault_path: Path to Obsidian vault
            include_daily_notes: Include daily notes
            include_templates: Include template files
        """
        super().__init__(**kwargs)
        self.vault_path = Path(vault_path).expanduser()
        self.include_daily_notes = include_daily_notes
        self.include_templates = include_templates

    async def iterate(self) -> AsyncIterator[DataPoint]:
        """Iterate over Obsidian notes."""
        if not self.vault_path.exists():
            return

        for md_file in self.vault_path.glob("**/*.md"):
            # Skip templates if not included
            if not self.include_templates and "template" in str(md_file).lower():
                continue

            # Skip .obsidian folder
            if ".obsidian" in str(md_file):
                continue

            try:
                content = md_file.read_text(encoding="utf-8")
            except Exception:
                continue

            if not content.strip():
                continue

            # Get file stats
            stat = md_file.stat()
            modified = datetime.fromtimestamp(stat.st_mtime)
            created = datetime.fromtimestamp(stat.st_ctime)

            # Extract frontmatter if present
            frontmatter = {}
            if content.startswith("---"):
                parts = content.split("---", 2)
                if len(parts) >= 3:
                    try:
                        import yaml
                        frontmatter = yaml.safe_load(parts[1]) or {}
                        content = parts[2].strip()
                    except Exception:
                        pass

            # Extract tags from content
            import re
            tags = re.findall(r"#(\w+)", content)
            tags = list(set(tags))

            # Extract links
            links = re.findall(r"\[\[([^\]]+)\]\]", content)

            # Determine note type
            note_path = md_file.relative_to(self.vault_path)
            is_daily = any(
                x in str(note_path).lower()
                for x in ["daily", "journal", "diary"]
            )

            if is_daily and not self.include_daily_notes:
                continue

            yield DataPoint(
                content=content,
                category=self.category,
                source_type=self.source_type,
                source_identifier=str(note_path),
                timestamp=modified,
                metadata={
                    "title": md_file.stem,
                    "path": str(note_path),
                    "created": created.isoformat(),
                    "modified": modified.isoformat(),
                    "frontmatter": frontmatter,
                    "links": links,
                    "is_daily_note": is_daily,
                },
                tags=["obsidian"] + tags,
            )


class NotionExportSource(DataSource):
    """Adapter for Notion workspace export."""

    source_type = "notion"
    category = DataCategory.NOTES

    def __init__(self, export_path: str, **kwargs):
        """
        Initialize Notion export source.

        Args:
            export_path: Path to Notion export (markdown or HTML)
        """
        super().__init__(**kwargs)
        self.export_path = Path(export_path).expanduser()

    async def iterate(self) -> AsyncIterator[DataPoint]:
        """Iterate over Notion pages."""
        if not self.export_path.exists():
            return

        # Process markdown files
        for md_file in self.export_path.glob("**/*.md"):
            try:
                content = md_file.read_text(encoding="utf-8")
            except Exception:
                continue

            if not content.strip():
                continue

            # Extract title from filename (Notion exports include page ID)
            filename = md_file.stem
            # Remove the Notion page ID suffix if present
            import re
            title_match = re.match(r"(.+?)\s+[a-f0-9]{32}$", filename)
            title = title_match.group(1) if title_match else filename

            # Get folder path for hierarchy
            rel_path = md_file.relative_to(self.export_path)
            hierarchy = list(rel_path.parent.parts)

            stat = md_file.stat()
            modified = datetime.fromtimestamp(stat.st_mtime)

            yield DataPoint(
                content=content,
                category=self.category,
                source_type=self.source_type,
                source_identifier=str(rel_path),
                timestamp=modified,
                metadata={
                    "title": title,
                    "path": str(rel_path),
                    "hierarchy": hierarchy,
                },
                tags=["notion"] + hierarchy,
            )


class AppleNotesSource(DataSource):
    """Adapter for Apple Notes database."""

    source_type = "apple_notes"
    category = DataCategory.NOTES

    def __init__(
        self,
        db_path: str = "~/Library/Group Containers/group.com.apple.notes/NoteStore.sqlite",
        **kwargs,
    ):
        """
        Initialize Apple Notes source.

        Args:
            db_path: Path to Notes database (requires Full Disk Access)
        """
        super().__init__(**kwargs)
        self.db_path = Path(db_path).expanduser()

    async def iterate(self) -> AsyncIterator[DataPoint]:
        """Iterate over Apple Notes."""
        if not self.db_path.exists():
            return

        try:
            conn = sqlite3.connect(str(self.db_path))
            conn.row_factory = sqlite3.Row

            # Query notes
            query = """
                SELECT
                    n.Z_PK as id,
                    n.ZTITLE1 as title,
                    n.ZSNIPPET as snippet,
                    n.ZMODIFICATIONDATE1 as modified,
                    n.ZCREATIONDATE1 as created,
                    nb.ZDATA as data
                FROM ZICCLOUDSYNCINGOBJECT n
                LEFT JOIN ZICNOTEDATA nb ON n.ZNOTEDATA = nb.Z_PK
                WHERE n.ZTITLE1 IS NOT NULL
                ORDER BY n.ZMODIFICATIONDATE1 DESC
            """

            cursor = conn.execute(query)

            for row in cursor:
                # Convert Apple's timestamp (seconds since 2001-01-01)
                modified = None
                if row["modified"]:
                    modified = datetime(2001, 1, 1) + \
                        __import__("datetime").timedelta(seconds=row["modified"])

                # Get note content
                content = row["snippet"] or ""

                # Try to extract text from data blob if available
                if row["data"] and not content:
                    try:
                        # Apple Notes stores content in a compressed format
                        import gzip
                        decompressed = gzip.decompress(row["data"])
                        content = decompressed.decode("utf-8", errors="ignore")
                    except Exception:
                        pass

                if not content:
                    continue

                yield DataPoint(
                    content=content,
                    category=self.category,
                    source_type=self.source_type,
                    source_identifier=f"note_{row['id']}",
                    timestamp=modified,
                    metadata={
                        "title": row["title"],
                    },
                    tags=["apple_notes"],
                )

            conn.close()

        except sqlite3.OperationalError as e:
            raise PermissionError(
                f"Cannot access Notes database. Grant Full Disk Access. Error: {e}"
            )


class EvernoteExportSource(DataSource):
    """Adapter for Evernote export (.enex files)."""

    source_type = "evernote"
    category = DataCategory.NOTES

    def __init__(self, export_path: str, **kwargs):
        """
        Initialize Evernote export source.

        Args:
            export_path: Path to .enex export file or folder
        """
        super().__init__(**kwargs)
        self.export_path = Path(export_path).expanduser()

    async def iterate(self) -> AsyncIterator[DataPoint]:
        """Iterate over Evernote notes."""
        import html
        import re
        import xml.etree.ElementTree as ET

        enex_files = []
        if self.export_path.is_file() and self.export_path.suffix == ".enex":
            enex_files = [self.export_path]
        elif self.export_path.is_dir():
            enex_files = list(self.export_path.glob("**/*.enex"))

        for enex_file in enex_files:
            try:
                tree = ET.parse(enex_file)
                root = tree.getroot()

                for note in root.findall(".//note"):
                    title = note.findtext("title", "")
                    content_elem = note.find("content")

                    # Parse content (ENML - Evernote Markup Language)
                    content = ""
                    if content_elem is not None and content_elem.text:
                        # Strip HTML/XML tags for plain text
                        text = content_elem.text
                        text = re.sub(r"<[^>]+>", " ", text)
                        text = html.unescape(text)
                        content = " ".join(text.split())

                    if not content and not title:
                        continue

                    # Parse dates
                    created = None
                    updated = None
                    created_str = note.findtext("created", "")
                    updated_str = note.findtext("updated", "")

                    for date_str, target in [(created_str, "created"), (updated_str, "updated")]:
                        if date_str:
                            try:
                                # Evernote date format: 20231215T120000Z
                                dt = datetime.strptime(date_str, "%Y%m%dT%H%M%SZ")
                                if target == "created":
                                    created = dt
                                else:
                                    updated = dt
                            except ValueError:
                                pass

                    # Extract tags
                    tags = [tag.text for tag in note.findall("tag") if tag.text]

                    yield DataPoint(
                        content=f"{title}\n\n{content}" if title else content,
                        category=self.category,
                        source_type=self.source_type,
                        source_identifier=f"evernote_{title}_{created_str}",
                        timestamp=updated or created,
                        metadata={
                            "title": title,
                            "created": created.isoformat() if created else None,
                            "notebook": enex_file.stem,
                        },
                        tags=["evernote"] + tags,
                    )

            except ET.ParseError:
                continue


class BearExportSource(DataSource):
    """Adapter for Bear notes export."""

    source_type = "bear"
    category = DataCategory.NOTES

    def __init__(self, export_path: str, **kwargs):
        """
        Initialize Bear export source.

        Args:
            export_path: Path to Bear export (markdown files)
        """
        super().__init__(**kwargs)
        self.export_path = Path(export_path).expanduser()

    async def iterate(self) -> AsyncIterator[DataPoint]:
        """Iterate over Bear notes."""
        if not self.export_path.exists():
            return

        for md_file in self.export_path.glob("**/*.md"):
            try:
                content = md_file.read_text(encoding="utf-8")
            except Exception:
                continue

            if not content.strip():
                continue

            # Bear uses # for tags in the content
            import re
            tags = re.findall(r"#([^\s#]+(?:/[^\s#]+)*)", content)
            tags = list(set(tags))

            stat = md_file.stat()
            modified = datetime.fromtimestamp(stat.st_mtime)

            yield DataPoint(
                content=content,
                category=self.category,
                source_type=self.source_type,
                source_identifier=str(md_file.name),
                timestamp=modified,
                metadata={
                    "title": md_file.stem,
                },
                tags=["bear"] + tags,
            )


class LocalDocumentsSource(DataSource):
    """Adapter for local documents (PDF, DOCX, TXT, etc.)."""

    source_type = "local_documents"

    def __init__(
        self,
        path: str,
        extensions: list[str] | None = None,
        recursive: bool = True,
        **kwargs,
    ):
        """
        Initialize local documents source.

        Args:
            path: Path to documents folder
            extensions: File extensions to include (default: common document types)
            recursive: Whether to search recursively
        """
        super().__init__(**kwargs)
        self.path = Path(path).expanduser()
        self.extensions = extensions or [".txt", ".md", ".pdf", ".docx", ".doc", ".rtf"]
        self.recursive = recursive

    @property
    def name(self) -> str:
        """Human-readable name of this source."""
        return f"Local Documents ({self.path.name})"

    @property
    def category(self) -> DataCategory:
        """Primary category of data from this source."""
        return DataCategory.DOCUMENTS

    async def is_available(self) -> bool:
        """Check if the path exists."""
        return self.path.exists()

    async def estimate_count(self) -> int:
        """Estimate number of documents."""
        if not self.path.exists():
            return 0
        pattern = "**/*" if self.recursive else "*"
        count = sum(1 for f in self.path.glob(pattern)
                   if f.is_file() and f.suffix.lower() in self.extensions)
        return count

    async def iterate(self) -> AsyncIterator[DataPoint]:
        """Iterate over local documents."""
        if not self.path.exists():
            return

        pattern = "**/*" if self.recursive else "*"

        for file_path in self.path.glob(pattern):
            if not file_path.is_file():
                continue

            if file_path.suffix.lower() not in self.extensions:
                continue

            try:
                content = await self._extract_content(file_path)
            except Exception:
                continue

            if not content or not content.strip():
                continue

            stat = file_path.stat()
            modified = datetime.fromtimestamp(stat.st_mtime)

            yield DataPoint(
                content=content,
                category=self.category,
                source_type=self.source_type,
                source_path=str(file_path),
                original_date=modified,
                raw_data={
                    "filename": file_path.name,
                    "extension": file_path.suffix,
                    "size_bytes": stat.st_size,
                    "path": str(file_path),
                },
                topics=["document", file_path.suffix.lstrip(".")],
            )

    async def _extract_content(self, file_path: Path) -> str:
        """Extract text content from a file."""
        suffix = file_path.suffix.lower()

        if suffix in [".txt", ".md"]:
            return file_path.read_text(encoding="utf-8", errors="ignore")

        elif suffix == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(file_path))
                text = []
                for page in reader.pages:
                    text.append(page.extract_text() or "")
                return "\n".join(text)
            except Exception:
                return ""

        elif suffix in [".docx", ".doc"]:
            try:
                from docx import Document
                doc = Document(str(file_path))
                return "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                return ""

        elif suffix == ".rtf":
            try:
                # Basic RTF parsing
                content = file_path.read_text(encoding="utf-8", errors="ignore")
                import re
                # Remove RTF control words
                content = re.sub(r"\\[a-z]+\d*\s?", "", content)
                content = re.sub(r"[{}]", "", content)
                return content
            except Exception:
                return ""

        return ""


class GoogleDocsExportSource(DataSource):
    """Adapter for Google Docs export (from Takeout)."""

    source_type = "google_docs"
    category = DataCategory.DOCUMENTS

    def __init__(self, export_path: str, **kwargs):
        """
        Initialize Google Docs export source.

        Args:
            export_path: Path to Google Takeout export folder
        """
        super().__init__(**kwargs)
        self.export_path = Path(export_path).expanduser()

    async def iterate(self) -> AsyncIterator[DataPoint]:
        """Iterate over Google Docs."""
        docs_path = self.export_path / "Drive"
        if not docs_path.exists():
            docs_path = self.export_path

        # Google Takeout exports docs as various formats
        for doc_file in docs_path.glob("**/*"):
            if not doc_file.is_file():
                continue

            suffix = doc_file.suffix.lower()
            if suffix not in [".docx", ".pdf", ".txt", ".md", ".html"]:
                continue

            try:
                content = await self._extract_content(doc_file)
            except Exception:
                continue

            if not content or not content.strip():
                continue

            stat = doc_file.stat()
            modified = datetime.fromtimestamp(stat.st_mtime)

            yield DataPoint(
                content=content,
                category=self.category,
                source_type=self.source_type,
                source_identifier=str(doc_file),
                timestamp=modified,
                metadata={
                    "title": doc_file.stem,
                    "format": suffix,
                },
                tags=["google_docs"],
            )

    async def _extract_content(self, file_path: Path) -> str:
        """Extract text content from a file."""
        suffix = file_path.suffix.lower()

        if suffix in [".txt", ".md"]:
            return file_path.read_text(encoding="utf-8", errors="ignore")

        elif suffix == ".docx":
            try:
                from docx import Document
                doc = Document(str(file_path))
                return "\n".join(p.text for p in doc.paragraphs)
            except Exception:
                return ""

        elif suffix == ".html":
            try:
                import re
                content = file_path.read_text(encoding="utf-8", errors="ignore")
                # Strip HTML tags
                content = re.sub(r"<[^>]+>", " ", content)
                return " ".join(content.split())
            except Exception:
                return ""

        elif suffix == ".pdf":
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(file_path))
                text = []
                for page in reader.pages:
                    text.append(page.extract_text() or "")
                return "\n".join(text)
            except Exception:
                return ""

        return ""
